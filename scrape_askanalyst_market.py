"""
AskAnalyst Market Intelligence Downloader
==========================================
Downloads market summaries, macro indicators, sector data, and PDF reports
from https://api.askanalyst.com.pk/api.

Output layout
─────────────
market_data/
  briefings/
    *.xlsx                   – PDF catalog indexes (title / date / URL)
    pdfs/                    – raw downloaded PDF files
    extracted/
      <title>.docx           – full text + tables in reading order (Word)
      <title>_tables.xlsx    – one worksheet per extracted table (Excel)
  news/
    psx_disclosures_latest.xlsx
  sectors/
    <sector>_metadata.json
    <sector>_<dataset>.json
    <sector>_<dataset>.xlsx  – flat export when data is tabular
  macro_indicators_index.xlsx

Dependencies
────────────
  pip install requests pandas openpyxl pdfplumber python-docx
"""

import io
import json
import os
import re
import time

import pandas as pd
import pdfplumber
import requests
from docx import Document


# ─── Config ───────────────────────────────────────────────────────────────────

BASE_API_URL          = "https://api.askanalyst.com.pk/api"
MARKET_DATA_DIR       = os.path.join(os.path.dirname(os.path.abspath(__file__)), "market_data")
BRIEFINGS_DIR         = os.path.join(MARKET_DATA_DIR, "briefings")
PDFS_DIR              = os.path.join(BRIEFINGS_DIR,   "pdfs")
EXTRACTED_DIR         = os.path.join(BRIEFINGS_DIR,   "extracted")
BRIEFINGS_CATALOGS_DIR = os.path.join(BRIEFINGS_DIR,   "catalogs")
BRIEFINGS_SUMMARY_DIR  = os.path.join(BRIEFINGS_DIR,   "summary")
SECTORS_DIR           = os.path.join(MARKET_DATA_DIR, "sectors")
NEWS_DIR              = os.path.join(MARKET_DATA_DIR, "news")
MACRO_DIR             = os.path.join(MARKET_DATA_DIR, "macroeconomics")

REQUEST_TIMEOUT = 30   # seconds — regular API calls
PDF_TIMEOUT     = 120  # seconds — PDF downloads can be large
DELAY           = 3    # seconds between every network call

_session = requests.Session()
_session.headers.update({
    "Accept":        "application/json, text/plain, */*",
    "Content-Type":  "application/json",
    "Authorization": "Bearer undefined",   # required by AskAnalyst API
    "Origin":        "https://www.askanalyst.com.pk",
    "Referer":       "https://www.askanalyst.com.pk/",
})

# Separate session for PDF downloads — no auth headers so Google Drive doesn't reject
_pdf_session = requests.Session()
_pdf_session.headers.update({
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/124.0 Safari/537.36",
})


# ─── Directory setup ──────────────────────────────────────────────────────────

def setup_dirs():
    for d in [BRIEFINGS_DIR, PDFS_DIR, EXTRACTED_DIR, SECTORS_DIR, NEWS_DIR, BRIEFINGS_CATALOGS_DIR, BRIEFINGS_SUMMARY_DIR, MACRO_DIR]:
        os.makedirs(d, exist_ok=True)
    for sector in ["cement", "fertilizer", "omc", "autos", "circulardebt"]:
        os.makedirs(os.path.join(SECTORS_DIR, sector), exist_ok=True)


# ─── Network helpers ──────────────────────────────────────────────────────────

_RETRY_DELAYS = [5, 15, 30]  # seconds to wait before each retry attempt


def api_get(path, params=None):
    url = f"{BASE_API_URL}/{path.lstrip('/')}"
    for attempt, backoff in enumerate([0] + _RETRY_DELAYS):
        if backoff:
            print(f"  [~] Waiting {backoff}s before retry {attempt}/{len(_RETRY_DELAYS)}...")
            time.sleep(backoff)
        try:
            r = _session.get(url, params=params, timeout=REQUEST_TIMEOUT)
            r.raise_for_status()
            return r.json()
        except requests.exceptions.HTTPError as e:
            print(f"  [-] HTTP error GET {path}: {e}")
            return None  # HTTP errors won't improve with retries
        except requests.exceptions.Timeout:
            print(f"  [-] Timeout GET {path} (attempt {attempt + 1})")
        except requests.exceptions.ConnectionError as e:
            print(f"  [-] Connection error GET {path} (attempt {attempt + 1}): {e}")
        except Exception as e:
            print(f"  [-] Error GET {path}: {e}")
            return None
    print(f"  [-] Giving up on GET {path} after {len(_RETRY_DELAYS)} retries.")
    return None


def api_post(path, payload):
    url = f"{BASE_API_URL}/{path.lstrip('/')}"
    for attempt, backoff in enumerate([0] + _RETRY_DELAYS):
        if backoff:
            print(f"  [~] Waiting {backoff}s before retry {attempt}/{len(_RETRY_DELAYS)}...")
            time.sleep(backoff)
        try:
            r = _session.post(url, json=payload, timeout=REQUEST_TIMEOUT)
            r.raise_for_status()
            return r.json()
        except requests.exceptions.HTTPError as e:
            print(f"  [-] HTTP error POST {path}: {e}")
            return None  # HTTP errors won't improve with retries
        except requests.exceptions.Timeout:
            print(f"  [-] Timeout POST {path} (attempt {attempt + 1})")
        except requests.exceptions.ConnectionError as e:
            print(f"  [-] Connection error POST {path} (attempt {attempt + 1}): {e}")
        except Exception as e:
            print(f"  [-] Error POST {path}: {e}")
            return None
    print(f"  [-] Giving up on POST {path} after {len(_RETRY_DELAYS)} retries.")
    return None


def download_pdf_bytes(url):
    """Download a PDF URL and return raw bytes, or None on failure."""
    try:
        # Use the plain PDF session — no AskAnalyst auth headers
        # which would cause Google Drive (and other hosts) to reject the request
        r = _pdf_session.get(url, timeout=PDF_TIMEOUT, allow_redirects=True)
        r.raise_for_status()
        return r.content
    except requests.exceptions.Timeout:
        print(f"  [-] Timeout downloading PDF: {url}")
    except Exception as e:
        print(f"  [-] Failed to download PDF {url}: {e}")
    return None


# ─── Save helpers ─────────────────────────────────────────────────────────────

def save_json(data, filepath):
    try:
        with open(filepath, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        print(f"  [OK] JSON  -> {filepath}")
    except Exception as e:
        print(f"  [-] Failed to save JSON {filepath}: {e}")


def save_excel(df, filepath, sheet_name="Sheet1"):
    try:
        with pd.ExcelWriter(filepath, engine="openpyxl") as writer:
            df.to_excel(writer, sheet_name=sheet_name[:31], index=False)
        print(f"  [OK] Excel ({len(df)} rows) -> {filepath}")
    except Exception as e:
        print(f"  [-] Failed to save Excel {filepath}: {e}")


def save_csv(df, filepath):
    try:
        df.to_csv(filepath, index=False, encoding="utf-8")
        print(f"  [OK] CSV   ({len(df)} rows) -> {filepath}")
    except Exception as e:
        print(f"  [-] Failed to save CSV {filepath}: {e}")


def save_txt(text, filepath):
    try:
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(text)
        print(f"  [OK] TXT   ({len(text)} chars) -> {filepath}")
    except Exception as e:
        print(f"  [-] Failed to save TXT {filepath}: {e}")


def sanitize_filename(name, max_len=80):
    """Strip characters that are illegal in filenames and truncate."""
    name = re.sub(r'[\\/*?:"<>|]', "", str(name))
    name = name.strip().replace(" ", "_")
    return name[:max_len] or "untitled"


# ─── Response normalisation ───────────────────────────────────────────────────

def flatten_to_df(data):
    """Best-effort conversion of an API response (list or dict) to a DataFrame."""
    if data is None:
        return None
    if isinstance(data, list):
        return pd.DataFrame(data) if data else None
    if isinstance(data, dict):
        for key in ("data", "results", "items", "records", "list", "posts"):
            if key in data and isinstance(data[key], list):
                return pd.DataFrame(data[key]) if data[key] else None
        return pd.DataFrame([data])
    return None


def normalise_pdf_df(df):
    """Rename common API column names to tidy Title / Date / URL columns."""
    rename = {}
    for col in df.columns:
        lc = col.lower()
        if lc in ("title", "name", "heading"):
            rename[col] = "Title"
        elif lc in ("date", "created_at", "createdat", "publish_date",
                    "published_at", "post_date"):
            rename[col] = "Date"
        elif lc in ("url", "link", "pdf", "pdf_url", "pdfurl",
                    "download_url", "file", "file_url"):
            rename[col] = "URL"
    return df.rename(columns=rename) if rename else df


# ─── PDF extraction ───────────────────────────────────────────────────────────

def _clean_table(raw_table):
    """Convert a pdfplumber raw table (list-of-lists with possible None) to clean strings."""
    return [
        [str(cell).strip() if cell is not None else "" for cell in row]
        for row in raw_table
        if any(cell is not None and str(cell).strip() for cell in row)
    ]


def _table_to_markdown(table):
    """Convert a clean table (list of lists) to a GitHub-flavoured markdown pipe table."""
    if not table:
        return ""
    def esc(cell):
        return cell.replace("|", "\\|").replace("\n", " ")
    headers = [esc(c) for c in table[0]]
    body_rows = []
    for row in table[1:]:
        padded = row + [""] * max(0, len(headers) - len(row))
        body_rows.append("| " + " | ".join(esc(c) for c in padded[:len(headers)]) + " |")
    return "\n".join([
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join(["---"] * len(headers)) + " |",
        *body_rows,
    ])


def extract_and_save_pdf(pdf_bytes, title, base_filename):
    """
    Extract all text and tables from a PDF.
    AI-readable  → <base_filename>.md           (markdown: headings + pipe tables)
    Human review → <base_filename>.docx         (Word: text + tables in reading order)
                   <base_filename>_tables.xlsx  (one worksheet per extracted table)
    """
    all_pages = []   # list of {"page": n, "text": str, "tables": [clean_table]}

    try:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                raw_tables = page.extract_tables() or []
                tables = [_clean_table(t) for t in raw_tables if t]
                text   = page.extract_text() or ""
                all_pages.append({"page": page_num, "text": text, "tables": tables})
    except Exception as e:
        print(f"  [-] pdfplumber could not read '{title}': {e}")
        return

    total_tables = sum(len(p["tables"]) for p in all_pages)
    total_chars  = sum(len(p["text"]) for p in all_pages)
    print(f"  [i] Extracted {total_tables} table(s), {total_chars} chars from '{title}'")

    # ── Excel: one sheet per table ────────────────────────────────────
    if total_tables:
        xl_path = os.path.join(EXTRACTED_DIR, base_filename + "_tables.xlsx")
        try:
            with pd.ExcelWriter(xl_path, engine="openpyxl") as writer:
                for page_info in all_pages:
                    for t_idx, table in enumerate(page_info["tables"], start=1):
                        if not table:
                            continue
                        headers = table[0]
                        rows    = table[1:] if len(table) > 1 else []
                        # Deduplicate header names (pdfplumber can emit blank/repeated headers)
                        seen, clean_headers = {}, []
                        for h in headers:
                            key = h or "Col"
                            if key in seen:
                                seen[key] += 1
                                key = f"{key}_{seen[key]}"
                            else:
                                seen[key] = 0
                            clean_headers.append(key)
                        df = pd.DataFrame(rows, columns=clean_headers)
                        sheet = f"P{page_info['page']}_T{t_idx}"[:31]
                        df.to_excel(writer, sheet_name=sheet, index=False)

                # "Full Text" sheet as a bonus
                text_rows = [
                    {"Page": p["page"], "Text": p["text"]}
                    for p in all_pages if p["text"].strip()
                ]
                if text_rows:
                    pd.DataFrame(text_rows).to_excel(writer, sheet_name="Full Text", index=False)

            print(f"  [OK] Excel tables -> {xl_path}")
        except Exception as e:
            print(f"  [-] Failed to save Excel for '{title}': {e}")

    # ── Word: human-readable version for review and verification ─────
    word_path = os.path.join(EXTRACTED_DIR, base_filename + ".docx")
    try:
        doc = Document()
        doc.add_heading(title, level=0)

        for page_info in all_pages:
            doc.add_heading(f"Page {page_info['page']}", level=2)

            # Text content (split into paragraphs)
            for line in page_info["text"].splitlines():
                line = line.strip()
                if line:
                    doc.add_paragraph(line)

            # Tables for this page
            for t_idx, table in enumerate(page_info["tables"], start=1):
                if not table:
                    continue
                doc.add_heading(f"Table {t_idx}", level=3)
                n_cols = max(len(row) for row in table)
                word_tbl = doc.add_table(rows=len(table), cols=n_cols)
                word_tbl.style = "Table Grid"
                for r_i, row in enumerate(table):
                    for c_i, cell_val in enumerate(row):
                        if c_i < n_cols:
                            cell = word_tbl.cell(r_i, c_i)
                            cell.text = cell_val
                            # Bold the header row
                            if r_i == 0:
                                for run in cell.paragraphs[0].runs:
                                    run.bold = True
                doc.add_paragraph()  # spacing after table

        doc.save(word_path)
        print(f"  [OK] Word doc  -> {word_path}")
    except Exception as e:
        print(f"  [-] Failed to save Word doc for '{title}': {e}")

    # ── Markdown: primary format for AI agent consumption ─────────────
    md_path = os.path.join(EXTRACTED_DIR, base_filename + ".md")
    try:
        lines = [f"# {title}", ""]
        for page_info in all_pages:
            lines.append(f"## Page {page_info['page']}")
            lines.append("")
            if page_info["text"].strip():
                lines.append(page_info["text"].strip())
                lines.append("")
            for t_idx, table in enumerate(page_info["tables"], start=1):
                if not table:
                    continue
                lines.append(f"### Table {t_idx}")
                lines.append("")
                lines.append(_table_to_markdown(table))
                lines.append("")
            lines.append("---")
            lines.append("")
        save_txt("\n".join(lines), md_path)
    except Exception as e:
        print(f"  [-] Failed to save Markdown for '{title}': {e}")


# ─── Section 1: General Market Summaries & PDF Catalogs ───────────────────────

def fetch_morning_briefing_chart():
    print("[*] Morning Briefing Chart (daily snapshot)...")
    data = api_get("/morningbriefingchart")
    time.sleep(DELAY)
    if data is None:
        print("  [-] Skipped.")
        return
    save_json(data, os.path.join(BRIEFINGS_SUMMARY_DIR, "morning_briefing_summary.json"))


def fetch_news(count=50):
    print(f"[*] PSX Latest News (last {count})...")
    data = api_get("/news/all", params={"page": 1, "postsperpage": count})
    time.sleep(DELAY)
    if data is None:
        print("  [-] Skipped.")
        return
    df = flatten_to_df(data)
    if df is None or df.empty:
        print("  [-] Skipped: could not parse response into a table.")
        return
    save_excel(df, os.path.join(NEWS_DIR, "latest_news.xlsx"), sheet_name="News")
    save_json(df.to_dict(orient="records"), os.path.join(NEWS_DIR, "latest_news.json"))
    save_csv(df, os.path.join(NEWS_DIR, "latest_news.csv"))


def fetch_pdf_catalog(endpoint, out_filename, max_pdfs=100):
    """
    1. Fetch the full JSON catalog → save as Excel index (all entries).
    2. Download, validate, and extract the most recent `max_pdfs` PDFs.
       Each PDF is saved raw to briefings/pdfs/ and its content extracted to
       briefings/extracted/ as both a Word doc and an Excel tables file.
    """
    print(f"[*] PDF catalog: /{endpoint}  (downloading latest {max_pdfs} PDFs)...")
    data = api_get(f"/{endpoint}")
    time.sleep(DELAY)
    if data is None:
        print("  [-] Skipped.")
        return

    df = flatten_to_df(data)
    if df is None or df.empty:
        print("  [-] Skipped: could not parse response into a table.")
        return

    df = normalise_pdf_df(df)

    # Save the full catalog index regardless of how many PDFs we download
    save_excel(df, os.path.join(BRIEFINGS_CATALOGS_DIR, out_filename), sheet_name="PDFs")
    print(f"  [i] Catalog has {len(df)} entries — downloading latest {min(max_pdfs, len(df))}.")

    url_col   = "URL"   if "URL"   in df.columns else None
    title_col = "Title" if "Title" in df.columns else None

    if url_col is None:
        print("  [!] No URL column found — skipping PDF downloads.")
        return

    downloaded = 0
    for _, row in df.iterrows():
        if downloaded >= max_pdfs:
            break

        url   = row.get(url_col)
        title = str(row.get(title_col, "untitled")) if title_col else "untitled"

        if not url or (isinstance(url, float) and pd.isna(url)):
            continue

        # Strip any trailing .pdf the API bakes into the title field
        title_clean = re.sub(r"\.pdf$", "", title, flags=re.IGNORECASE).strip()

        print(f"  [{downloaded + 1}/{max_pdfs}] {title_clean[:70]}...")
        pdf_bytes = download_pdf_bytes(str(url))
        time.sleep(DELAY)

        if pdf_bytes is None:
            continue

        # Validate: must start with %PDF and be a reasonable size
        if len(pdf_bytes) < 5000:
            print(f"  [!] Skipping — only {len(pdf_bytes)} bytes (not a valid PDF).")
            continue
        if not pdf_bytes.startswith(b"%PDF"):
            print(f"  [!] Skipping — no PDF header (server returned an error page).")
            continue

        # Save raw PDF
        raw_path = os.path.join(PDFS_DIR, sanitize_filename(title_clean) + ".pdf")
        try:
            with open(raw_path, "wb") as f:
                f.write(pdf_bytes)
        except Exception as e:
            print(f"  [-] Could not save raw PDF: {e}")

        extract_and_save_pdf(pdf_bytes, title_clean, sanitize_filename(title_clean))
        time.sleep(DELAY)
        downloaded += 1

    print(f"  [OK] {downloaded} PDF(s) downloaded and extracted.")


# ─── Section 2: Macroeconomic Indicators Index ────────────────────────────────

def fetch_macro_indicators():
    print("[*] Macroeconomic Indicators Index (/searchlist)...")
    data = api_get("/searchlist")
    time.sleep(DELAY)
    if data is None:
        print("  [-] Skipped.")
        return
    df = flatten_to_df(data)
    if df is None or df.empty:
        print("  [-] Skipped: could not parse response.")
        return
    save_excel(df, os.path.join(MACRO_DIR, "macro_indicators_index.xlsx"),
               sheet_name="Indicators")


# ─── Section 3: Sector-Specific Fundamentals ──────────────────────────────────

SECTOR_POST_CONFIGS = {
    "cement": [
        {
            "label": "domestic_monthly_dispatches",
            "frequency": "monthly",
            "payload_extra": {
                "indices": "dispatches",
                "type": "domestic",
                "parameter": "value",
                "growth": "value"
            }
        },
        {
            "label": "export_monthly_dispatches",
            "frequency": "monthly",
            "payload_extra": {
                "indices": "dispatches",
                "type": "export",
                "parameter": "value",
                "growth": "value"
            }
        }
    ],
    "fertilizer": [
        {
            "label": "urea_monthly_sales",
            "frequency": "monthly",
            "payload_extra": {
                "type": "sales",
                "product": "Urea",
                "parameter": "local_sales",
                "growth": "value"
            }
        },
        {
            "label": "urea_monthly_production",
            "frequency": "monthly",
            "payload_extra": {
                "type": "production",
                "product": "Urea",
                "parameter": "production",
                "growth": "value"
            }
        }
    ],
    "omc": [
        {
            "label": "motor_spirit_monthly_sales",
            "frequency": "monthly",
            "payload_extra": {
                "salestab": "productwise",
                "product": "Motor Spirit",
                "parameter": "sales",
                "growth": "value"
            }
        },
        {
            "label": "hsd_monthly_sales",
            "frequency": "monthly",
            "payload_extra": {
                "salestab": "productwise",
                "product": "High Speed Diesel",
                "parameter": "sales",
                "growth": "value"
            }
        }
    ],
    "autos": [
        {
            "label": "passenger_car_monthly_sales",
            "frequency": "monthly",
            "payload_extra": {
                "salestab": "sales",
                "category": "All Products",
                "product": "All Products",
                "parameter": "sales",
                "growth": "value"
            }
        },
        {
            "label": "passenger_car_monthly_production",
            "frequency": "monthly",
            "payload_extra": {
                "salestab": "production",
                "category": "All Products",
                "product": "All Products",
                "parameter": "production",
                "growth": "value"
            }
        }
    ],
    "circulardebt": [
        {
            "label": "circular_debt_monthly_flows",
            "frequency": "monthly",
            "payload_extra": {
                "type": "flow",
                "parameter": "value",
                "growth": "value"
            }
        }
    ]
}


def fetch_sector(sector):
    print(f"\n[+] Sector: {sector.upper()}")
    endpoint = f"/sector/{sector}"

    # Step 1 – GET baseline metadata
    print(f"  [*] GET {endpoint} (metadata)...")
    metadata = api_get(endpoint)
    time.sleep(DELAY)
    if metadata is None:
        print(f"  [-] No metadata returned. Skipping {sector}.")
        return
    save_json(metadata, os.path.join(SECTORS_DIR, sector, f"{sector}_metadata.json"))

    # Step 2 – POST for each configured dataset
    configs = SECTOR_POST_CONFIGS.get(sector, [])
    if not configs:
        print(f"  [!] No POST configs defined for {sector}. Metadata only.")
        return

    for cfg in configs:
        label     = cfg["label"]
        frequency = cfg["frequency"]

        # Pull full date range from metadata for this frequency (handling nested sector structures)
        dates_dict = metadata.get("dates", {})
        if sector.lower() == "autos" and "Autos" in dates_dict:
            dates_list = dates_dict["Autos"].get(frequency, [])
        elif sector.lower() == "omc" and "Omc" in dates_dict:
            dates_list = dates_dict["Omc"].get(frequency, [])
        else:
            dates_list = dates_dict.get(frequency, [])

        if dates_list:
            edate = dates_list[0]["start_date"]
            # Limit range to latest 6 periods to avoid server-side timeouts on historical data
            sdate = dates_list[min(5, len(dates_list)-1)]["start_date"]
        else:
            sdate = "2026-01-01"
            edate = "2026-06-01"

        payload = {
            "sdate":     sdate,
            "edate":     edate,
            "frequency": frequency,
            **cfg["payload_extra"]
        }
        if sector != "circulardebt":
            payload["company"] = "All Companies"

        print(f"  [*] POST {endpoint} -> {label} ({sdate} to {edate})...")
        result = api_post(endpoint, payload)
        time.sleep(DELAY)

        if result is None:
            print(f"  [-] No data for {label}.")
            continue

        save_json(result, os.path.join(SECTORS_DIR, sector, f"{sector}_{label}.json"))

        df = flatten_to_df(result)
        if df is not None and not df.empty:
            save_excel(df, os.path.join(SECTORS_DIR, sector, f"{sector}_{label}.xlsx"),
                       sheet_name=label[:31])
            save_csv(df, os.path.join(SECTORS_DIR, sector, f"{sector}_{label}.csv"))


# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    print("=" * 60)
    print("  AskAnalyst Market Intelligence Downloader")
    print("=" * 60)

    setup_dirs()

    print("\n=== Section 1: Market Summaries & PDFs ===")
    fetch_morning_briefing_chart()
    fetch_news(count=50)
    fetch_pdf_catalog("morningbriefing",   "morning_briefing_pdfs.xlsx",       max_pdfs=3)
    fetch_pdf_catalog("technicalresearch", "technical_research_pdfs.xlsx",     max_pdfs=3)
    fetch_pdf_catalog("marketroundup",     "market_roundup_pdfs.xlsx",         max_pdfs=3)
    fetch_pdf_catalog("reports/1000",      "research_reports_catalog.xlsx",    max_pdfs=50)

    print("\n=== Section 2: Macroeconomic Indicators ===")
    fetch_macro_indicators()

    print("\n=== Section 3: Sector Fundamentals ===")
    for sector in ["cement", "fertilizer", "omc", "autos", "circulardebt"]:
        fetch_sector(sector)

    print("\n" + "=" * 60)
    print("  Download complete.")
    print("=" * 60)


if __name__ == "__main__":
    main()
